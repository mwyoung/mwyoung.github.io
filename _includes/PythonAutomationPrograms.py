#01_intro.py
#!/usr/bin/python3
# Says hello, asks for name
print('Hello')

myStr = input('Enter text: ')
print('text is: ' + myStr)
print('length:', end='')
print(len(myStr))

print('Type in a number: ', end='')
inputNum = input()
print('Num++ is ' + str(int(inputNum)+1) + '!')

#==========================================================================================
#02_for_while.py
#flow
#!/usr/bin/python3

name = 'Abc'
if name == 'Abc':
    print('sup')
else:
    print(':(')
print('done')

quit()

#str
#!/usr/bin/python3
print('Enter text')
name = input()
if name:
    print('entered text')
else:
    print(':(')
print('done')

quit()

#while
#!/usr/bin/python3
spam = 0
while spam < 5:
    print('text')
    spam = spam + 1

quit()

#whiletext
#!/usr/bin/python3
text = ''
while text != 'text':
    print('type text to quit')
    text = input()
    #or can do:
    #if text == 'text'
    #   break
print('done')
quit()

#whilecontinue
#!/usr/bin/python3
spam = 0
while spam < 5:
    spam = spam + 1
    if spam == 3:
        continue    # goes back to start, does not print
    print('text')

quit()

#for
#!/usr/bin/python3
print('text')
for i in range(5):
    print('loop '+ str(i) )
quit()

#forsum
#!/usr/bin/python3
total = 0
for i in range(100 + 1):
    total = total + i
print(total)
quit()

#==========================================================================================
#03_rand_func_var_exit_cp.py
#rand
#!/usr/bin/python3
#!/usr/bin/python3
import random
print(str(random.randint(1,10)))
quit()

#exit
#!/usr/bin/python3
import sys
print('here')
sys.exit()
print('never here')

#copypaste
#!/usr/bin/python3
#need to install xclip and python3-pyqt5...
import pyperclip
pyperclip.copy('test')
print('clipboard is:' + pyperclip.paste())
quit()

#func
#!/usr/bin/python3

def printtest(str):
    print('func called with string: ' + str)

printtest('abc')
printtest('')       #cannot be blank
printtest('test')

quit()

#one
#!/usr/bin/python3

def plusone(num):
    return num + 1

newNumber = plusone(5)
print(newNumber)

quit()

#global local
#!/usr/bin/python3

def spam():
    global var  #needed or else variable doesn't change
    print('in function')
    var = 'test'

var = 'abc'
print(var)
spam()
print(var)

quit()

#==========================================================================================
#04_try_except.py
#try
#!/usr/bin/python3

def div42by(divideBy):
    try:
        return 42 / divideBy
    except ZeroDivisionError:
        print('Err: no div 0')

print(div42by(2))
print(div42by(12))
print(div42by(0))
print(div42by(1))

quit()

#except
#!/usr/bin/python3

print('Enter a number')
num = input()
try:
    if int(num) >= 4:
        print('is > 4')
    elif int(num) > 0:
        print('not a lot')
    else:
        print('is neg')
except ValueError:
    print('not a number')

quit()

#==========================================================================================
#05_rand_num.py
#!/usr/bin/python3
#This is a guess the number game
import random   #for randint function

print('Enter a name:')
name = input()
print(name +', guesss a number between 1 and 20 (including)')
secretNumber = random.randint(1,20)

#print('Debug: num is ' + str(secretNumber))

for guessesTaken in range(1,7): #6 iterations
    print('Take a guess')
    guess = int(input()) #convert to a integer

    if guess < secretNumber:
        print('too low')
    elif guess > secretNumber:
        print('too high')
    else: #guess correct
        break

if guess == secretNumber:
    print('found number in ' + str(guessesTaken) + ' guesses')
else:
    print('The number was ' + str(secretNumber))

quit()

#==========================================================================================
#06_list.py
#for list
#!/usr/bin/python3

array = ['a', 'b', 'c', 'd']

for i in range(len(array)):
    print('Index ' + str(i) + ' is ' + array[i])

quit()

#list mutable
#!/usr/bin/python3

def eggs(cheese):
    cheese.append('hello')

spam = [1, 2, 3]
print(spam)
eggs(spam)
print(spam)

quit()

#==========================================================================================
#07_char_count.py
#!/usr/bin/python3
import pprint #pretty printing

def count(msg):
    count = {} #count number of characters exist

    for character in msg.upper():
        count.setdefault(character,0) #set character to 0, if not exist
        count[character] = count[character] + 1

    #print(count)
    #pprint.pprint(count) # in order, single column
    rjtext = pprint.pformat(count) # in order, single column
    print(rjtext)

#msg = 'It was a bright cold day in April, and the clocks were stricking thirteen.'

#multiline quote
# from https://automatetheboringstuff.com/files/rj.txt
msg = '''
The Project Gutenberg EBook of Romeo and Juliet, by William Shakespeare

This eBook is for the use of anyone anywhere at no cost and with
almost no restrictions whatsoever.  You may copy it, give it away or
re-use it under the terms of the Project Gutenberg License included
with this eBook or online at www.gutenberg.org/license


Title: Romeo and Juliet

Author: William Shakespeare

Posting Date: May 25, 2012 [EBook #1112]
Release Date: November, 1997  [Etext #1112]

Language: English
'''

count(msg)

quit()

#==========================================================================================
#08_str.py
#!/usr/bin/python3

str = 'this\nis a test for using \' chars'
str2 = r'literal\n'
print(str)
print(str2)

print(str.lower())
print(str.upper())
str4 = str.replace('\n',' ')
print(str4)

quit()

#==========================================================================================
#09_launch.py
#!/usr/bin/python3
import sys

print('Hello world')

print('args: ', end='')
print(sys.argv)     #command arguments

quit()

#==========================================================================================
#10_regex_pipe.py
#phone
#!/usr/bin/python3

#takes in number, format 123-456-7890
def isPhoneNum(text):
    if len(text) != 12:
        return False #not phone-num size
    for i in range(0,3):
        if not text[i].isdecimal():
            return False #no area code
    if text[3] != '-':
        return False #missing dash
    for i in range(4,7):
        if not text[i].isdecimal():
            return False #no 3 digits
    if text[7] != '-':
        return False #missing dash
    for i in range(8,12):
        if not text[i].isdecimal():
            return False #no 4 digits
    return True

print(isPhoneNum('415-555-1234'))
print(isPhoneNum('15-555-1234'))

#more code if number is in text
def testIPN(message):
    foundNum = False
    for i in range(len(message)):
        chunk = message[i:i+12]     #slice, not worry if going past bounds
        if isPhoneNum(chunk):
            print('Phone number ' + chunk + ' in ' + str(i) + ' loops')
            foundNum = True
    if not foundNum:
        print('Could not find any phone number')

message = 'call me at 415-555-1011 tomorrow, or at 415-555-9999'
print(message)
testIPN(message)
message = 'call me at 415-55-1011 tomorrow, or at 415-555-999'
print(message)
testIPN(message)

quit()

#phone regex
#!/usr/bin/python3
import re

message = 'call me at 415-555-1011 tomorrow, or at 415-555-9999'
# \d = digit
phoneNumRegex = re.compile(r'\d\d\d-\d\d\d-\d\d\d\d')
mo = phoneNumRegex.search(message) #return match option, first occurance
print(mo.group()) #finds first match
mo = phoneNumRegex.findall(message) #return list of matches
print(mo) #finds first match

quit()

#phone group
#!/usr/bin/python3
import re

message = 'my number is 415-555-4242'
phoneNumRegex = re.compile(r'(\d\d\d)-(\d\d\d-\d\d\d\d)') # \d = digit. () = group
mo = phoneNumRegex.search(message) #return match option, first occurance
print(mo.group() + ', 1: ' + mo.group(1) + ' 2: ' + mo.group(2))

message = 'my number is (415)-555-4242'
phoneNumRegex = re.compile(r'\(\d\d\d\)-\d\d\d-\d\d\d\d') # \d = digit. () = group
mo = phoneNumRegex.search(message) #return match option, first occurance
print(mo.group())

quit()

#pipe
#!/usr/bin/python3

import re

batRegex = re.compile(r'Bat(man|mobile|copter|sub)') #group, seperate with pipe
mo = batRegex.search('Batmobile lost a wheel')
print(mo.group() + ' found in group ' + mo.group(1))
mo = batRegex.search('Batmotorcycle lost a wheel')
print(mo.group()) #return none, crash

quit()

#repeat
#!/usr/bin/python3

import re

#question?
batRegex = re.compile(r'Bat(wo)?man') #group appear 0 or 1 times
print(batRegex.search('The Adventures of Batman'))
print(batRegex.search('The Adventures of Batwoman'))

#phone numbs
phoneRegex = re.compile(r'(\d\d\d-)?\d\d\d-\d\d\d\d') #does not require area code
print(phoneRegex.search('my num is 555-1234'))
print(phoneRegex.search('my num is 123-555-1234'))

#asterisk*
batRegex = re.compile(r'Bat(wo)*man') #group appear 0+ times
print(batRegex.search('The Adventures of Batman'))
print(batRegex.search('The Adventures of Batwowoman'))

#add+
batRegex = re.compile(r'Bat(wo)+man') #group appear 1+ times
print(batRegex.search('The Adventures of Batwoman'))

#escape\
regex = re.compile(r'\+\*\?')
print(regex.search('learn about +*? regex'))
regex = re.compile(r'(\+\*\?)+') #appear 1+ times (in order)
print(regex.search('learn about +*?+*?+*? regex'))

#exact match times
haRegex = re.compile(r'(Ha){3}')
print(haRegex.search('he "HaHaHa"'))
haRegex = re.compile(r'(Ha){3,5}')#including, match first 5 only
print(haRegex.search('he "HaHaHaHaHa"'))
#match phone number, optional area code, 3 in a row, optinal ending ,
phRegex = re.compile(r'((\d\d\d-)?\d\d\d-\d\d\d\d(,)?){3}')
print(phRegex.search('nums are 415-555-1234,555-4242,212-555-0000'))

#testing
digitRegex = re.compile(r'(\d){3,5}') #match digits, greedy
print(digitRegex.search('1234567890'))
digitRegex = re.compile(r'(\d){3,5}?') #match digits, non-greedy
print(digitRegex.search('1234567890'))

quit()

#findall
#!/usr/bin/python3
import re
msg = '123-456-7890, 342-234-2315, 556-234-7152, and 234-6151'
phoneRegex = re.compile(r'\d\d\d-\d\d\d-\d\d\d\d')
print(phoneRegex.search(msg)) #find first
print(phoneRegex.findall(msg)) #list of values

phoneRegex = re.compile(r'(\d\d\d)-(\d\d\d-\d\d\d\d)') #2 groups
print(phoneRegex.findall(msg)) #list of values

lyrics = '12 drum, 11 piper, 10 lords, 9 dance, 8 milk, 7 swan, 6 geese, ' + \
        '5 rings, 4 birds, 3 hens, 2 doves, 1 tree'
# 1 or more digits, space (or use ' '), 1+ word char
xmas = re.compile(r'\d+\s\w+')
print(xmas.findall(lyrics))

vowelRegex = re.compile(r'[aeiouAEIOU]') #4'(a|e|i|o|u)'
print(vowelRegex.findall('Robocop eats baby food.'))
vowelRegex = re.compile(r'[aeiouAEIOU]{2}') #read 2 vowels
print(vowelRegex.findall('Robocop eats baby food.'))
vowelRegex = re.compile(r'[^aeiouAEIOU]') #opposite
print(vowelRegex.findall('Robocop eats baby food.'))
vowelRegex = re.compile(r'[^aeiouAEIOU\s]') #opposite
print(vowelRegex.findall('Robocop eats baby food.'))

quit()

#dot star caret dollar
#!/usr/bin/python3
import re
beginHello = re.compile(r'^Hello') #begins with hello
print(beginHello.search('Hello there!'))
endWorld = re.compile(r'world!$') #$ends with world
print(endWorld.search('Hello world!'))

allDigit = re.compile(r'^\d+$') #one or more digits, no chars
print(allDigit.search('31241324512431'))

atRegex = re.compile(r'.at')
print(atRegex.findall('The cat in the hat sat on the flat mat'))
atRegex = re.compile(r'.{1,2}at') #at preciding 1 or 2 chars, any
print(atRegex.findall('The cat in the hat sat on the flat mat'))

str = 'First Name: Test Last Name: User'
nameRegex = re.compile(r'First Name: (.*) Last Name: (.*)') #use groups
print(nameRegex.findall(str))

str2 = '<To serve humans> for dinner.>'
nongreedy = re.compile(r'<(.*?)>') #non greedy, 1st match
print(nongreedy.findall(str2))
greedy = re.compile(r'<(.*)>') #non greedy, 1st match
print(greedy.findall(str2))

str3 = 'line one\nline two\nline three'
dotstar = re.compile(r'.*')
print(dotstar.search(str3)) #match until newline
dotstar = re.compile(r'.*', re.DOTALL)
print(dotstar.search(str3)) #match until newline

vowelRegex = re.compile(r'[aeiou]', re.IGNORECASE) #4'(a|e|i|o|u)'
print(vowelRegex.findall('A robocop eats baby food.'))

quit()

#sub verbose
#!/usr/bin/python3
import re
str='Agent Alice gave docs to Agent Bob'
name = re.compile(r'Agent \w+')
print(name.findall(str))
print(name.sub('REDACTED',str))
name = re.compile(r'Agent (\w)\w*') #groups, get first letter of group, 0+ more other chars
print(name.findall(str))
print(name.sub(r'Agent \1****',str)) #replace word in group

verb = re.compile(r'''
(\d\d\d-)|    #area code
(\(\d\d\d\))  #area code with paren, no dash
\d\d\d    #digits
-
\d\d\d\d
\sx\d{2,4}  #extension, like x1234
''', re.VERBOSE)

quit()

#scraper
#!/usr/bin/python3
import re, pyperclip

# regex for phone num
phoneRegex = re.compile(r'''
# types: 415-123-1234, 123-1234, (415) 123-1234, 123-1234 ext 12345, ext. 12345, x12345
(                           #one group, no touples if using findall
((\d\d\d)|(\(\d\d\d\)))?    # area code (optional)
(\s|-)    # first separator
\d\d\d    # 3 digits
-       # separator
\d\d\d\d    # 4 digits
(((ext(\.)?\s)|x)
  (\d{2,5}))?    # extension (optional)
)
''', re.VERBOSE)

# regex for email
emailRegex = re.compile(r'''
#some.+_thing@something.com
[a-zA-Z0-9_.+]+    #name, custom
@                #@
[a-zA-Z0-9_.+]+    #domain
''', re.VERBOSE)

# get text off clipboard
text = pyperclip.paste()

# extract email/phone from text
extractedPhone = phoneRegex.findall(text)
# copy extract email/phone to clipboard
extractedEmail = emailRegex.findall(text)

allPhoneNumbers = [] #empty list
for phoneNum in extractedPhone:
    allPhoneNumbers.append(phoneNum[0])

#print(extractedPhone) #testing
#print(allPhoneNumbers)
#print(extractedEmail)

# copy email/phone to clipboard
#one line, newline seperated
results = '\n'.join(allPhoneNumbers) + '\n' + '\n'.join(extractedEmail)
pyperclip.copy(results);
#print(results)

quit()

#==========================================================================================
#11_file.py
#file
#!/usr/bin/python3
import os

print(os.path.join('f1','f2','f3','f.png'))
print(os.sep)
print(os.getcwd())
strPath = os.getcwd() + os.sep + 'test.ext'
print(strPath)
print(os.path.dirname(strPath))
print(os.path.basename(strPath))

quit()

#filesize
#!/usr/bin/python3
import os

totalSize = 0
filePath = os.getcwd()
for filename in os.listdir(filePath):
    if not os.path.isfile(os.path.join(filePath, filename)):
        continue
    totalSize = totalSize + os.path.getsize(os.path.join(filePath, filename))

print('File sizes are ' + str(totalSize) + ' bytes')

quit()

#rw
#!/usr/bin/python3
import os

filePath = os.getcwd() + os.sep + '11_hello.txt'
if not os.path.isfile(filePath):
    print('File does not exist')
    quit()

helloFile = open(filePath)
print(helloFile.read())
helloFile.seek(0) #go to start of file
print(helloFile.readlines())
helloFile.close()

baconFile = open('bacon.txt', 'w') #overwrite
baconBytes = baconFile.write('Bacon is not a vegetable.')
print('bytes written: ' + str(baconBytes))
baconFile.close()

baconFile = open('bacon.txt', 'a') #append
baconBytes = baconFile.write('\n\nBacon is a thing')
print('bytes written: ' + str(baconBytes))
baconFile.close()

baconFile = open('bacon.txt', 'r') #read
print(baconFile.read())
baconFile.close()

quit()


#vars
#!/usr/bin/python3
import os, shelve
shelfFile = shelve.open('shelvedata')
shelfFile['things'] = ['a','b','c','d']
shelfFile.close()

shelfFile = shelve.open('shelvedata')
print(shelfFile['things']) #key/value
print(list(shelfFile.keys()))
print(list(shelfFile.values()))

shelfFile.close()
quit()

#delete
#!/usr/bin/python3
import os

os.chdir(os.path.expanduser('~') + os.sep + 'Pictures')
print('location: ' + os.getcwd())

for filename in os.listdir():
    if filename.endswith('.jpg'):
        #os.unlink(filenm) #uncomment and change filenm to rm
        print(filename)

quit()

#walk
#!/usr/bin/python3
import os

folderPath = os.getcwd() + os.sep + 'folder'
print('folder: ' + folderPath)

for folderName, subfolders, filenames in os.walk(folderPath):
    print('Folder: ' + folderName)
    print('subfolders in ' + folderName + ':' + str(subfolders))
    print('filenames in ' + folderName + ':' + str(filenames))
    print(===) #newline

    ''' text/comment
    for subfolder in subfolders:
        if 'text' in subfolder:
            #os.rmdir(subfolder)
    for file in filenames:
        if file.endswith('.py'):
            #shutil.copy(os.path.join(folderName, file), \
            #        os.path.join(folderName, file + '.bk')) #rename
    '''

quit()

#==========================================================================================
#12_debug_log_assert.py
#raise
#!/usr/bin/python3

"""

***************
*             *
*             *
*             *
***************

"""

def boxPrint(symbol, width, height):
    if len(symbol) != 1:
        raise Exception('"symbol" needs to be a string of length 1')
    if (width < 2) or (height < 2):
        raise Exception('"width" and "height" must be >= 2')

    print(symbol * width)   #use string replication
    for i in range(height-2):
        print(symbol + (' ' * (width-2)) + symbol)
    print(symbol * width)

boxPrint('*', 15, 5)
boxPrint('o', 5, 7)

#boxPrint('**', 5, 10) #error
#boxPrint('*', 1, 1) #error

quit()

#traceback
#!/usr/bin/python3
import traceback
try:
    raise Exception('error msg')
except:
    errorFile = open('traceback_error_log.txt', 'a') #add to existing file
    errorFile.write(traceback.format_exc())
    errorFile.close()
    print('The traceback info was written to 12_traceback_error_log.txt')

quit()

#assert
#!/usr/bin/python3

market_2nd = {'ns': 'green', 'ew':'red'} #north-south = green, east-west = red

def switchLights(intersection):
    for key in intersection.keys(): #loop over keys
        if intersection[key] == 'green':
            intersection[key] = 'yellow'
        elif intersection[key] == 'yellow':
            intersection[key] = 'red'
        elif intersection[key] == 'red':
            intersection[key] = 'green'
    assert 'red' in intersection.values(), 'neither light is red' + str(intersection)

print(market_2nd)
switchLights(market_2nd)
print(market_2nd)

quit()

#log
#!/usr/bin/python3
#buggy factorial
import logging
#logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logging.basicConfig(level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',datefmt='%H:%M:%S')
#logging.disable(logging.CRITICAL) #turn off logging
logging.debug('start of program')

def factorial(n):
    logging.debug('start of factorial(%s)' % (n))
    total = 1
    #for i in range (n+1): #is buggy
    for i in range (1, n + 1):
        total *= i
        logging.debug('i is %s, total is %s' % (i, total))

    logging.debug('return value is %s' % (total))
    return total

print(factorial(5))

logging.debug('end of program')
quit()


#debug
#!/usr/bin/python3

print('1st num:',end=' ')
first = input()
print('2nd num:',end=' ')
second = input()
print('3rd num:',end=' ')
third = input()
print('The sum is ' + first + second + third)

quit()

#breakpoint
#!/usr/bin/python3
import random

heads = 0

for i in range(1,1001): #1000 coin flips
    if random.randint(0,1) == 1:
        heads = heads + 1
    if i == 500:
        print('1/2 done')

print('heads: ' + str(heads) + ' times')

quit()

#==========================================================================================
#13_web.py
#webbrowser
#!/usr/bin/python3
#open google maps with clipboard or sysargs
import webbrowser, sys, pyperclip
#webbrowser.open('https://automatetheboringstuff.com') #opens browser

#check if command line args passed
if len(sys.argv) > 1:
    # sys.argv = ['program', 'arg1', 'arg2', ...]]
    address = ' '.join(sys.argv[1:]) #join list, seperate w/ space, avoid [0]
else:
    address = pyperclip.paste()

#https://www.google.com/maps/place/<address>
webbrowser.open('https://www.google.com/maps/place/' + address)

quit()

#request
#!/usr/bin/python3
import requests
res = requests.get('https://automatetheboringstuff.com/files/rj.txt')
print('response: ' + str(res.status_code))
print('len: ' + str(len(res.text)))

'''
playFile = open('rj.txt', 'wb')
for chunk in res.iter_content(100000): #return content in bytes, iterate content len
    playFile.write(chunk)
playFile.close()
'''

badRes = requests.get('https://automatetheboringstuff.com/ljdsaflkjsfdalkj') #badurl
print('response: ' + str(badRes.status_code))
badRes.raise_for_status()

quit()

#parse
#!/usr/bin/python3
import bs4, requests #beautifulsoup4, downloading
#get price info:
def getAmazonPrice(productURL):
	#need to use user-agent for Amazon, due to bot detection
    header = {
        'User-Agent': '...',
    }
    res = requests.get(productURL, headers=header)
    res.raise_for_status() #raise exception if error
    #print('response: ' + str(res.status_code))

    soup = bs4.BeautifulSoup(res.text, 'html.parser')
    #css selector to find, highlight + right click + Inspect -> right click -> copy -> copy selector (css path)
    #elems = soup.select('#mediaNoAccordion > div.a-row > div.a-column.a-span4.a-text-right.a-span-last > span.a-size-medium.a-color-price.header-price') #address of price
    #elems = soup.select('#mediaNoAccordion .header-price')
    elems = soup.select('#content > div:nth-child(2) > div.column.column-block.small-12.medium-3.medium-text-right > div:nth-child(1) > div.column.small-8.medium-12 > p > span > span')
    return elems[0].text.strip()

#has lots of scraper blockers
#getAmazonPrice('https://www.amazon.com/Automate-Boring-Stuff-Python-Programming-ebook/dp/B00WJ049VU')
price = getAmazonPrice('https://camelcamelcamel.com/Automate-Boring-Stuff-Python-Programming/product/1593275994')
print('price: ' + price)

quit()

#selenium
#!/usr/bin/python3
#requires firefox
from selenium import webdriver

browser = webdriver.Firefox() #launch, store in object, opens new window
browser.get('https://automatetheboringstuff.com')
elem = browser.find_element_by_css_selector('.entry-content > ol:nth=child(15) > li:nth-child(1) > a:nth-child(1)')
elem.click() #click on link
elems = browser.find_elements_by_css_selector('p')
len(elems)

#search
searchElem = browser.find_element_by_css_selector('.search-field')
searchElem.send_keys('test')

#read
elem = browser.find_element_by_css_selector('entry-content > p:nth-child(4)')
print(elem.text)

quit()

#==========================================================================================
#14_xls_pdf_doc.py
#excel
#!/usr/bin/python3
import openpyxl, os
#os.chdir('\doc\location')

#read xlsx - from http://autbor.com/example.xlsx
workbook = openpyxl.load_workbook('example.xlsx')
print(type(workbook))
#print(workbook.get_sheet_names()) #depreciated
#sheet = workbook.get_sheet_by_name('Sheet1') #depreciated
print(workbook.sheetnames)
sheet = workbook['Sheet1']
print(type(sheet))
print(str(sheet['A1'].value))
print(str(sheet.cell(row=1,column=2).value))
for i in range(1,8):
    print(i, sheet.cell(row=i, column=2).value, end='; ')
print('')

print('------------')
#write
wb = openpyxl.Workbook() #create
print(wb.sheetnames)
st = wb[wb.sheetnames[0]]
st['A1'] = 42
st['A2'] = 'Hello'
#wb.save('file.xlsx')
st2 = wb.create_sheet()
print(wb.sheetnames)
st2.title = 'new title'
print(wb.sheetnames)
st3 = wb.create_sheet(index=0, title='test2')
print(wb.sheetnames)

quit()

#pdf
#!/usr/bin/python3
import PyPDF2

#read pdf - from http://autbor.com/meetingminutes1.pdf
pdfFile = open('pdf1.pdf', 'rb')
reader = PyPDF2.PdfFileReader(pdfFile)
print('pages: ' + str(reader.numPages))
page = reader.getPage(0)
print(page.extractText()[:10])
#for pageNum in range(reader.numPages):
#   print(reader.getPage(pageNum).extractText())

#modificaiton
pdfFile2 = open('pdf2.pdf', 'rb')
reader2 = PyPDF2.PdfFileReader(pdfFile2)

writer = PyPDF2.PdfFileWriter()

for pageNum in range(reader.numPages):
    page = reader.getPage(pageNum)
    writer.addPage(page)

for pageNum in range(reader2.numPages):
    page = reader2.getPage(pageNum)
    writer.addPage(page)

outputFile = open('pdf3.pdf', 'wb')
writer.write(outputFile)
outputFile.close()
pdfFile.close()
pdfFile.close()

quit()

#doc
#!/usr/bin/python3
import docx

#read docx - from http://autbor.com/demo.docx
doc = docx.Document('demo.docx')
print(doc.paragraphs[0].text)

par = doc.paragraphs[1]
print(par.text)
print(par.runs)
print(par.runs[0].text)
#d.save('test.docx') #saves file

# new doc
doc2 = docx.Document()
doc2.add_paragraph('this is a paragraph, ')
doc2.add_paragraph('this is another paragraph')
par2 = doc2.paragraphs[0]
par2.add_run('this is a new run')
par2.runs[1].bold = True
#d.save('test4.docx')

#get doc as string
def getText(filename):
    doc3 = docx.Document(filename)
    fullText = [] #blank list
    for para in doc3.paragraphs:
        fullText.append(para.text) #append text of each object
    return '\n'.join(fullText) #single string

print(getText('demo.docx'))

quit()

#==========================================================================================
#15_email.py
#email
#!/usr/bin/python3
import smtplib
conn = smtplib.SMTP('smtp.gmail.com', 587) #connection object, port number 587
conn.ehlo() #start connection
conn.starttls() #secure
conn.login('email@gmail.com', 'password')
conn.sendmail('from@email', 'to@email', 'Subject: subject here\n\nEmail goes here\n\ntext')
conn.quit()

quit()

#check email
#!/usr/bin/python3
import imapclient, pyzmail
conn = imapclient.IMAPClient('imap.gmail.com', ssl=True) #depends on email provider
conn.login(email, pswd)
conn.select_folder('INBOX', readonly=True) #not delete email
# all, before, on, since, subject, body, text, etc
conn.search(['SINCE 20-AUG-2019']) #imap syntax
rawMsg = conn.fetch([num], ['BODY[]', 'FLAGS'])
pyzmail.PyzMessage.factory(rawMsg[num][b'BODY[]'])
print(msg.get_subject())
print(msg.get_addresses('from')
if msg.text_part != None:
    print(msg.text_part.get_payload().decode('UTF-8'))
#conn.list_folders()

quit()

#==========================================================================================
#16_mouse.py
#!/usr/bin/python3
import pyautogui

print('Screen size: ' + str(pyautogui.size()))
width, height = pyautogui.size()
print('Current location: ' + str(pyautogui.position()))

pyautogui.moveTo(600,600, duration=0.1) #move to x,y
print('Current location: ' + str(pyautogui.position()))
pyautogui.moveRel(120,120) #move to x,y
print('Current location: ' + str(pyautogui.position()))
#pyautogui.click(x, y)

quit()

